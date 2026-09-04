"""Developer-only offline QA harness for paired DOCX/PDF reports.

The command deliberately writes only to an explicit output directory. It does
not update a catalogue, approval store, or any source report. It is not part of
the CarbonFactorResolver production runtime and is not exposed by the CFR API.
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import importlib.metadata
import json
import re
import subprocess
from collections.abc import Iterable, Mapping
from dataclasses import asdict, dataclass, replace
from datetime import UTC, datetime
from decimal import Decimal
from pathlib import Path
from typing import Any

from a1_factor_engine.adapters import HttpCatalogFactorRepository
from a1_factor_engine.engine import A1FactorResolutionEngine
from a1_factor_engine.material_registry import DEFAULT_MATERIAL_REGISTRY
from a1_factor_engine.models import ResolutionStatus

EXPECTED_STAGES = ("A1", "A2", "A3", "TOTAL")
PARSER_VERSION = "cfr.true-data-extractor/v0.13.1"
REPORT_NAME_PATTERN = re.compile(r"^(?P<report_id>\d{2})--")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def canonical_json_bytes(value: object) -> bytes:
    return json.dumps(
        value, ensure_ascii=False, sort_keys=True, separators=(",", ":")
    ).encode("utf-8")


def parse_number(value: str) -> Decimal:
    return Decimal(value.replace(",", "").strip())


def label_value(paragraphs: Iterable[str], label: str) -> str:
    for paragraph in paragraphs:
        for line in paragraph.splitlines():
            if label in line:
                return line.split(label, 1)[1].strip()
    raise ValueError(f"required label not found: {label}")


def normalized_token(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.casefold())


def first_year(value: str) -> int:
    match = re.search(r"\b(20\d{2})\b", value)
    if not match:
        raise ValueError(f"year not found: {value}")
    return int(match.group(1))


def infer_product_name_cn(path: Path) -> str:
    """Extract the product label from the report filename without a case registry."""
    tail = re.sub(r"^\d+--", "", path.stem)
    if "--" in tail:
        tail = tail.rsplit("--", 1)[1]
    elif "_" in tail:
        tail = tail.split("_", 1)[1]
    tail = tail.lstrip("_-")
    suffixes = (
        "产品碳足迹报告说明_证书边框",
        "_中英双语_证书边框_new",
        "中英双语_证书边框_new",
        "_中英双语_证书边框",
        "中英双语_证书边框",
        "报告说明_证书边框",
        "_证书边框",
    )
    for suffix in suffixes:
        if tail.endswith(suffix):
            tail = tail[: -len(suffix)]
            break
    return re.sub(r"\s+", "", tail).strip("_-")


@dataclass(frozen=True)
class SourcePair:
    report_id: str
    product_name_cn: str
    docx_path: str
    pdf_path: str
    docx_sha256: str
    pdf_sha256: str
    docx_size: int
    pdf_size: int


@dataclass(frozen=True)
class ExtractedFactor:
    factor_id: str
    report_id: str
    factor_name_cn: str
    factor_name_en: str
    material_name_cn: str
    category: str
    stage: str
    value: float
    unit: str
    source: str
    source_version: str
    source_year: int
    boundary: str
    boundary_modules: tuple[str, ...]
    applicability: str
    certificate_no: str
    accounting_period: str
    docx_path: str
    docx_sha256: str
    docx_evidence: str
    pdf_path: str
    pdf_sha256: str
    pdf_evidence: str
    cross_format_verified: bool
    product_name_en: str = ""
    value_raw: str = ""
    display_precision: int = 2
    pdf_table_index: int = 0
    pdf_row_index: int = 0
    pdf_column_index: int = 4
    pdf_cell_bbox: tuple[float, float, float, float] | None = None
    parser_version: str = PARSER_VERSION
    extraction_confidence: float = 1.0
    license: str = "internal-read-only-source-document"
    source_quality_status: str = "VERIFIED"
    admission_eligible: bool = True


@dataclass(frozen=True)
class QualityFinding:
    report_id: str
    severity: str
    code: str
    message: str
    evidence: str


def _indexed_reports(paths: Iterable[Path], suffix: str) -> dict[str, Path]:
    grouped: dict[str, list[Path]] = {}
    for path in paths:
        match = REPORT_NAME_PATTERN.match(path.name)
        if not match:
            raise ValueError(f"invalid {suffix} report filename (expected NN--...): {path.name}")
        grouped.setdefault(match.group("report_id"), []).append(path)
    duplicates = {key: values for key, values in grouped.items() if len(values) != 1}
    if duplicates:
        details = {key: sorted(path.name for path in values) for key, values in duplicates.items()}
        raise ValueError(f"duplicate {suffix} report IDs: {details}")
    return {key: values[0] for key, values in grouped.items()}


def source_pairs(source_dir: Path, *, expected_pairs: int | None = None) -> tuple[SourcePair, ...]:
    docx_by_id = _indexed_reports(source_dir.glob("*.docx"), "DOCX")
    pdf_by_id = _indexed_reports(source_dir.glob("*.pdf"), "PDF")
    if set(docx_by_id) != set(pdf_by_id):
        raise ValueError(
            "DOCX/PDF report IDs differ: "
            f"docx_only={sorted(set(docx_by_id) - set(pdf_by_id))}, "
            f"pdf_only={sorted(set(pdf_by_id) - set(docx_by_id))}"
        )
    if expected_pairs is not None and len(docx_by_id) != expected_pairs:
        raise ValueError(f"expected {expected_pairs} paired reports, found {len(docx_by_id)}")

    return tuple(
        SourcePair(
            report_id=report_id,
            product_name_cn=infer_product_name_cn(docx_by_id[report_id]),
            docx_path=str(docx_by_id[report_id].resolve()),
            pdf_path=str(pdf_by_id[report_id].resolve()),
            docx_sha256=sha256_file(docx_by_id[report_id]),
            pdf_sha256=sha256_file(pdf_by_id[report_id]),
            docx_size=docx_by_id[report_id].stat().st_size,
            pdf_size=pdf_by_id[report_id].stat().st_size,
        )
        for report_id in sorted(docx_by_id)
    )


def _pdf_lifecycle_cells(page: object) -> dict[str, tuple[str, int, int, tuple[float, float, float, float] | None]]:
    matches: dict[str, tuple[str, int, int, tuple[float, float, float, float] | None]] = {}
    tables = page.find_tables()  # type: ignore[attr-defined]
    for table_index, table in enumerate(tables):
        rows = table.extract()
        if not rows or len(rows[0]) < 5 or "carbonfootprint" not in normalized_token(str(rows[0][4])):
            continue
        for row_index, row in enumerate(rows[1:], 1):
            if len(row) < 5:
                continue
            row_label = normalized_token(str(row[0] or ""))
            stage = next((value for value in EXPECTED_STAGES[:-1] if value.casefold() in row_label), None)
            if stage is None and "total" in row_label:
                stage = "TOTAL"
            if stage is None:
                continue
            bbox = None
            try:
                bbox = table.rows[row_index].cells[4]
            except (AttributeError, IndexError):
                pass
            matches[stage] = (str(row[4] or "").strip(), table_index, row_index, bbox)
    return matches


def extract_pair(pair: SourcePair) -> tuple[tuple[ExtractedFactor, ...], tuple[QualityFinding, ...]]:
    try:
        import pdfplumber
        from docx import Document
    except ImportError as exc:  # pragma: no cover - environment diagnostic
        raise RuntimeError(
            "true-data extraction requires the optional acceptance-tools dependencies"
        ) from exc

    document = Document(pair.docx_path)
    paragraphs = tuple(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    product_en = label_value(paragraphs, "Product Name:")
    specification = label_value(paragraphs, "Specification/Model:")
    certificate_no = label_value(paragraphs, "Certificate No.:")
    accounting_period = label_value(paragraphs, "Accounting Period:")
    boundary_text = label_value(paragraphs, "System Boundary:")
    total_label = label_value(paragraphs, "Product Carbon Footprint per Unit:")
    total_value = parse_number(re.search(r"[\d,.]+", total_label).group(0))  # type: ignore[union-attr]
    year = first_year(accounting_period)
    issue_date = label_value(paragraphs, "Issue Date:")
    source_version = f"{certificate_no}; issued {issue_date}"
    source = label_value(paragraphs, "Evaluation and Reporting Organization:")

    if len(document.tables) != 1 or len(document.tables[0].rows) != 5:
        raise ValueError(
            f"{pair.report_id}: expected one five-row lifecycle table, "
            f"found {len(document.tables)} tables"
        )
    rows = [
        tuple(cell.text.replace("\n", " ").strip() for cell in row.cells)
        for row in document.tables[0].rows[1:]
    ]
    observed_stages = tuple(row[0].split()[0] if row[0].startswith("A") else "TOTAL" for row in rows)
    if observed_stages != EXPECTED_STAGES:
        raise ValueError(f"{pair.report_id}: unexpected lifecycle rows {observed_stages!r}")

    with pdfplumber.open(pair.pdf_path) as pdf:
        pdf_pages = tuple(page.extract_text() or "" for page in pdf.pages)
        pdf_cells = _pdf_lifecycle_cells(pdf.pages[1]) if len(pdf.pages) >= 2 else {}
    if len(pdf_pages) != 2:
        raise ValueError(f"{pair.report_id}: expected two PDF pages, found {len(pdf_pages)}")
    compact_pdf = normalized_token("\n".join(pdf_pages))
    product_verified = normalized_token(product_en) in compact_pdf

    factors: list[ExtractedFactor] = []
    findings: list[QualityFinding] = []
    stage_values: list[Decimal] = []
    stage_net_emissions: list[Decimal] = []
    for row_number, (stage, row) in enumerate(zip(EXPECTED_STAGES, rows, strict=True), 2):
        net_emissions = Decimal(row[3].replace(",", ""))
        value = Decimal(row[4].replace(",", ""))
        if stage != "TOTAL":
            stage_values.append(value)
            stage_net_emissions.append(net_emissions)
        page = 2
        if stage == "TOTAL":
            category = "product_carbon_footprint"
            boundary = "cradle-to-gate"
            modules: tuple[str, ...] = ("A1", "A2", "A3")
            applicability = (
                f"{product_en}; specification {specification}; per tonne of product; "
                f"{boundary_text}; accounting period {accounting_period}"
            )
            name_cn = f"{pair.product_name_cn}产品碳足迹"
            name_en = f"{product_en} product carbon footprint"
        else:
            category = "lifecycle_stage_carbon_footprint"
            boundary = stage
            modules = (stage,)
            subprocess = row[2]
            applicability = (
                f"{product_en}; lifecycle stage {stage}; {subprocess}; per tonne of product; "
                f"accounting period {accounting_period}"
            )
            name_cn = f"{pair.product_name_cn} {stage}阶段碳足迹"
            name_en = f"{product_en} {stage} lifecycle-stage carbon footprint"
        pdf_cell = pdf_cells.get(stage)
        numeric_verified = bool(pdf_cell and parse_number(pdf_cell[0]) == value)
        factors.append(
            ExtractedFactor(
                factor_id=f"TD-{pair.report_id}-{stage}",
                report_id=pair.report_id,
                factor_name_cn=name_cn,
                factor_name_en=name_en,
                material_name_cn=pair.product_name_cn,
                category=category,
                stage=stage,
                value=float(value),
                unit="kgCO2e/t",
                source=source,
                source_version=source_version,
                source_year=year,
                boundary=boundary,
                boundary_modules=modules,
                applicability=applicability,
                certificate_no=certificate_no,
                accounting_period=accounting_period,
                docx_path=Path(pair.docx_path).name,
                docx_sha256=pair.docx_sha256,
                docx_evidence=(
                    "paragraph containing Product Carbon Footprint per Unit; "
                    f"Table 1 row {row_number}, Carbon Footprint column"
                    if stage == "TOTAL"
                    else f"Table 1 row {row_number}, Carbon Footprint column"
                ),
                pdf_path=Path(pair.pdf_path).name,
                pdf_sha256=pair.pdf_sha256,
                pdf_evidence=(
                    f"page {page}, table {pdf_cell[1] if pdf_cell else 'missing'}, "
                    f"row {pdf_cell[2] if pdf_cell else 'missing'}, Carbon Footprint column"
                ),
                cross_format_verified=bool(product_verified and numeric_verified),
                product_name_en=product_en,
                value_raw=row[4].replace(",", "").strip(),
                display_precision=len(row[4].rsplit(".", 1)[1]) if "." in row[4] else 0,
                pdf_table_index=pdf_cell[1] if pdf_cell else -1,
                pdf_row_index=pdf_cell[2] if pdf_cell else -1,
                pdf_cell_bbox=pdf_cell[3] if pdf_cell else None,
            )
        )

    total_row_value = Decimal(factors[-1].value_raw)
    if abs(sum(stage_values) - total_row_value) > Decimal("0.02"):
        findings.append(
            QualityFinding(
                pair.report_id,
                "error",
                "FOOTPRINT_STAGE_SUM_MISMATCH",
                f"A1+A2+A3={sum(stage_values)} but total={total_row_value}",
                "DOCX Table 1 rows 2-5; PDF page 2 lifecycle table",
            )
        )
    total_net = Decimal(rows[-1][3].replace(",", ""))
    if abs(sum(stage_net_emissions) - total_net) > Decimal("0.02"):
        findings.append(
            QualityFinding(
                pair.report_id,
                "error",
                "NET_EMISSIONS_SUM_MISMATCH",
                f"A1+A2+A3={sum(stage_net_emissions)} tCO2 but total={total_net} tCO2",
                "DOCX Table 1 rows 2-5; PDF page 2 lifecycle table",
            )
        )
    if abs(total_value - total_row_value) > Decimal("0.02"):
        findings.append(
            QualityFinding(
                pair.report_id,
                "error",
                "PAGE_TOTAL_MISMATCH",
                f"page 1 total={total_value} but page 2 total={total_row_value}",
                "DOCX page-1 total paragraph and Table 1 total row; PDF pages 1-2",
            )
        )
    for factor in factors:
        if not factor.cross_format_verified:
            findings.append(
                QualityFinding(
                    pair.report_id,
                    "error",
                    "CROSS_FORMAT_TEXT_CHECK_FAILED",
                    f"DOCX factor {factor.factor_id} was not found verbatim in normalized PDF text",
                    factor.pdf_evidence,
                )
            )
    rejected = any(finding.severity == "error" for finding in findings)
    status = "REJECTED" if rejected else "VERIFIED"
    return (
        tuple(replace(factor, source_quality_status=status, admission_eligible=not rejected) for factor in factors),
        tuple(findings),
    )


def catalog_payload(factors: tuple[ExtractedFactor, ...]) -> dict[str, Any]:
    records = []
    for factor in factors:
        base_name = factor.product_name_en
        records.append(
            {
                "record_id": factor.factor_id,
                "category": "epd_indicator",
                "subject_type": "finished_product",
                "code": factor.factor_id,
                "name": base_name,
                "aliases": [factor.material_name_cn],
                "primary_value": factor.value,
                "value_raw": factor.value_raw,
                "display_precision": factor.display_precision,
                "primary_unit": factor.unit,
                "source": factor.source,
                "source_id": factor.certificate_no,
                "source_version": factor.source_version,
                "document_status": "PUBLISHED",
                "source_type": "epd",
                "year": factor.source_year,
                "boundary": factor.boundary,
                "boundary_modules": list(factor.boundary_modules),
                "indicator": "GWP-total",
                "declared_product": base_name,
                "source_quality_status": factor.source_quality_status,
                "admission_eligible": factor.admission_eligible,
                "cross_format_verified": factor.cross_format_verified,
                "parser_version": factor.parser_version,
                "extraction_confidence": factor.extraction_confidence,
                "license": factor.license,
                "scope": factor.applicability,
                "source_document_locator": f"evidence://report/{factor.report_id}/pdf/{factor.pdf_sha256}",
                "source_document_sha256": factor.pdf_sha256,
                "page": "2",
                "table": "Lifecycle Process",
                "row": factor.stage,
                "evidence_cell_bbox": factor.pdf_cell_bbox,
                "notes": factor.pdf_evidence,
            }
        )
    record_sha = hashlib.sha256(canonical_json_bytes(records)).hexdigest()
    return {
        "catalog_version": "true-data-read-only-snapshot/v1",
        "database": {"name": "read-only-report-extract", "sha256": record_sha},
        "records": records,
    }


def ingestion_cases(factors: tuple[ExtractedFactor, ...]) -> list[dict[str, Any]]:
    all_ids = {factor.factor_id for factor in factors}
    cases: list[dict[str, Any]] = []
    for factor in factors:
        base_name = factor.product_name_en
        product_variant_required = "product_variant" in (
            DEFAULT_MATERIAL_REGISTRY.resolve(base_name).identity.unresolved_attributes
        )
        acceptable = [factor.factor_id] if factor.admission_eligible and not product_variant_required else []
        cases.append(
            {
                "schema_version": "cfr-true-data-acceptance/v1",
                "case_id": f"CASE-{factor.factor_id}",
                "case_type": (
                    "source_quality_control" if not factor.admission_eligible
                    else "ambiguity_control" if product_variant_required
                    else "extracted_factor"
                ),
                "factor_id": factor.factor_id,
                "correct_material_entity": base_name,
                "request": {
                    "material_name": base_name,
                    "subject_type": "finished_product",
                    "quantity": 1.0,
                    "quantity_unit": "t",
                    "year": factor.source_year,
                    "boundary": factor.boundary,
                    "target_factor_unit": factor.unit,
                    "top_k": 5,
                },
                "acceptable_candidates": acceptable,
                "forbidden_candidates": sorted(all_ids - set(acceptable)),
                "expected_more_input": product_variant_required,
                "expected_abstention": not factor.admission_eligible,
                "answer_basis": (
                    "Exact declared product, accounting year, lifecycle boundary, unit, "
                    "and paired DOCX/PDF evidence must agree."
                ),
            }
        )

    controls = (
        ("CTRL-SLIDING-GATE", "Sliding Gate", True),
        ("CTRL-SILICA-BRICK", "Silica Brick", True),
        ("CTRL-PRECAST", "Precast Shape", True),
        ("CTRL-WEAR-CASTABLE", "Wear-Resistant Castable", True),
        ("CTRL-ALUMINA-RAW", "Alumina raw material", False),
        ("CTRL-ALUMINIUM-METAL", "Primary aluminium metal", False),
        ("CTRL-STEEL-FIBRE-RAW", "Steel fibre raw material", False),
        ("CTRL-NATURAL-GAS", "Natural gas", False),
    )
    for case_id, material_name, more_input in controls:
        cases.append(
            {
                "schema_version": "cfr-true-data-acceptance/v1",
                "case_id": case_id,
                "case_type": "ambiguity_control" if more_input else "abstention_control",
                "factor_id": None,
                "correct_material_entity": material_name,
                "request": {
                    "material_name": material_name,
                    "subject_type": "finished_product" if more_input else "raw_material",
                    "quantity": 1.0,
                    "quantity_unit": "t",
                    "boundary": "cradle-to-gate",
                    "target_factor_unit": "kgCO2e/t",
                    "top_k": 5,
                },
                "acceptable_candidates": [],
                "forbidden_candidates": sorted(all_ids),
                "expected_more_input": more_input,
                "expected_abstention": not more_input,
                "answer_basis": (
                    "Ambiguous family label requires clarification."
                    if more_input
                    else "The extracted report snapshot contains no exact raw-material factor."
                ),
            }
        )
    return cases


def write_json(path: Path, value: object) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_jsonl(path: Path, values: Iterable[Mapping[str, Any]]) -> None:
    with path.open("w", encoding="utf-8", newline="\n") as stream:
        for value in values:
            stream.write(json.dumps(value, ensure_ascii=False, sort_keys=True) + "\n")


def load_frozen_cases(path: Path) -> list[dict[str, Any]]:
    cases = [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
    if not all(isinstance(case, dict) for case in cases):
        raise ValueError("holdout manifest lines must be JSON objects")
    case_ids = [str(case.get("case_id") or "") for case in cases]
    if not all(case_ids) or len(case_ids) != len(set(case_ids)):
        raise ValueError("holdout case_id values must be non-empty and unique")
    return cases


def write_factor_csvs(output_dir: Path, factors: tuple[ExtractedFactor, ...], cases: list[dict[str, Any]]) -> None:
    import csv

    cases_by_factor = {case["factor_id"]: case for case in cases if case["factor_id"]}
    target = output_dir / "factor_test_tables"
    target.mkdir(parents=True, exist_ok=True)
    for factor in factors:
        case = cases_by_factor[factor.factor_id]
        rows = [
            ("factor_id", factor.factor_id),
            ("factor_name_cn", factor.factor_name_cn),
            ("factor_name_en", factor.factor_name_en),
            ("category", factor.category),
            ("value", factor.value),
            ("unit", factor.unit),
            ("source", factor.source),
            ("source_version", factor.source_version),
            ("source_year", factor.source_year),
            ("boundary", factor.boundary),
            ("applicability", factor.applicability),
            ("correct_material_entity", case["correct_material_entity"]),
            ("acceptable_candidates", "|".join(case["acceptable_candidates"])),
            ("forbidden_candidate_rule", "any candidate not listed as acceptable"),
            ("expected_more_input", case["expected_more_input"]),
            ("expected_abstention", case["expected_abstention"]),
            ("docx_evidence", f"{factor.docx_path} :: {factor.docx_evidence}"),
            ("pdf_evidence", f"{factor.pdf_path} :: {factor.pdf_evidence}"),
        ]
        with (target / f"{factor.factor_id}.csv").open("w", encoding="utf-8-sig", newline="") as stream:
            writer = csv.writer(stream)
            writer.writerow(("field", "value"))
            writer.writerows(rows)


def evidence_coverage(candidate: Any) -> float:
    source = candidate.source
    fields = (
        source.source_id,
        source.provider,
        source.locator,
        source.factor_unit,
        source.factor_kind.value,
        source.indicator,
        source.declared_product,
        source.boundary,
        source.source_document_sha256,
        source.page,
        source.table,
        source.row,
        source.metadata.get("cross_format_verified"),
        source.metadata.get("parser_version"),
        source.metadata.get("extraction_confidence"),
        source.metadata.get("license"),
        source.metadata.get("source_quality_status"),
    )
    return sum(value is not None and str(value).strip() != "" for value in fields) / len(fields)


def more_input_is_correct(status: str, expected_more_input: bool) -> bool:
    if status == ResolutionStatus.ERROR.value:
        return False
    return (status == ResolutionStatus.MORE_INPUT_NEEDED.value) == expected_more_input


def aggregate_acceptance_metrics(
    results: list[dict[str, Any]],
    *,
    preset_sha_before: str,
    preset_sha_after: str,
    catalog_record_anchor: str,
) -> dict[str, Any]:
    retrieval = [result for result in results if result["recall_at_5"] is not None]
    stage_retrieval = [
        result for result in retrieval
        if str(result.get("factor_id") or "").rsplit("-", 1)[-1] in {"A1", "A2", "A3"}
    ]
    candidate_count = sum(int(result["candidate_count"]) for result in results)
    wrong_count = sum(int(result["wrong_candidate_count"]) for result in results)
    abstentions = [result for result in results if result["abstention_correct"] is not None]
    positive_more_input = [result for result in results if result["expected_more_input"]]
    negative_more_input = [result for result in results if not result["expected_more_input"]]
    evidence_sum = sum(float(result["evidence_coverage_sum"]) for result in results)
    evidence_count = sum(int(result["evidence_candidate_count"]) for result in results)
    return {
        "case_count": len(results),
        "extracted_factor_case_count": len(retrieval),
        "recall_at_5": (
            sum(bool(result["recall_at_5"]) for result in retrieval) / len(retrieval)
            if retrieval
            else None
        ),
        "top_1_accuracy": (
            sum(bool(result.get("recall_at_1", result["recall_at_5"])) for result in retrieval) / len(retrieval)
            if retrieval else None
        ),
        "wrong_candidate_rate": wrong_count / candidate_count if candidate_count else 0.0,
        "qualified_candidate_precision": 1 - (wrong_count / candidate_count) if candidate_count else 1.0,
        "exact_stage_wrong_candidate_rate": (
            sum(int(result["wrong_candidate_count"]) for result in stage_retrieval)
            / sum(int(result["candidate_count"]) for result in stage_retrieval)
            if sum(int(result["candidate_count"]) for result in stage_retrieval) else 0.0
        ),
        "correct_abstention_rate": (
            sum(bool(result["abstention_correct"]) for result in abstentions) / len(abstentions)
            if abstentions
            else None
        ),
        "abstention_case_count": len(abstentions),
        "more_input_reasonableness_rate": (
            sum(bool(result["more_input_correct"]) for result in results) / len(results)
            if results
            else None
        ),
        "more_input_positive_recall": (
            sum(result["observed_status"] == ResolutionStatus.MORE_INPUT_NEEDED.value for result in positive_more_input)
            / len(positive_more_input) if positive_more_input else None
        ),
        "more_input_positive_case_count": len(positive_more_input),
        "more_input_negative_specificity": (
            sum(result["observed_status"] != ResolutionStatus.MORE_INPUT_NEEDED.value for result in negative_more_input)
            / len(negative_more_input) if negative_more_input else None
        ),
        "more_input_negative_case_count": len(negative_more_input),
        "unnecessary_question_rate": (
            sum(result["observed_status"] == ResolutionStatus.MORE_INPUT_NEEDED.value for result in negative_more_input)
            / len(negative_more_input) if negative_more_input else None
        ),
        "evidence_completeness_rate": evidence_sum / evidence_count if evidence_count else None,
        "evidence_metadata_presence_rate": evidence_sum / evidence_count if evidence_count else None,
        "wrong_candidate_count": wrong_count,
        "returned_candidate_count": candidate_count,
        "case_error_count": sum(result["observed_status"] == ResolutionStatus.ERROR.value for result in results),
        "preset_sha256_before": preset_sha_before,
        "preset_sha256_after": preset_sha_after,
        "catalog_snapshot_sha256": catalog_record_anchor,
    }


def release_gate(metrics: Mapping[str, Any], holdout_metrics: Mapping[str, Any] | None) -> dict[str, Any]:
    checks = {
        "ingestion_recall_at_5": metrics.get("recall_at_5") == 1.0,
        "ingestion_wrong_candidate_rate": float(metrics.get("wrong_candidate_rate", 1.0)) == 0.0,
        "ingestion_correct_abstention": float(metrics.get("correct_abstention_rate", 0.0)) >= 0.95,
        "ingestion_more_input_positive": float(metrics.get("more_input_positive_recall", 0.0)) >= 0.95,
        "ingestion_more_input_negative": float(metrics.get("more_input_negative_specificity", 0.0)) >= 0.95,
        "ingestion_evidence": metrics.get("evidence_completeness_rate") == 1.0,
        "ingestion_no_errors": int(metrics.get("case_error_count", 1)) == 0,
        "holdout_present": holdout_metrics is not None,
    }
    if holdout_metrics is not None:
        checks.update({
            "holdout_recall_at_5": holdout_metrics.get("recall_at_5") == 1.0,
            "holdout_wrong_candidate_rate": float(holdout_metrics.get("wrong_candidate_rate", 1.0)) <= 0.05,
            "holdout_correct_abstention": float(holdout_metrics.get("correct_abstention_rate", 0.0)) >= 0.95,
            "holdout_negative_sample_size": int(holdout_metrics.get("abstention_case_count", 0)) >= 20,
            "holdout_more_input_positive": float(holdout_metrics.get("more_input_positive_recall", 0.0)) >= 0.95,
            "holdout_more_input_negative": float(holdout_metrics.get("more_input_negative_specificity", 0.0)) >= 0.95,
            "holdout_evidence": holdout_metrics.get("evidence_completeness_rate") == 1.0,
            "holdout_no_errors": int(holdout_metrics.get("case_error_count", 1)) == 0,
        })
    return {"passed": all(checks.values()), "checks": checks}


def write_acceptance_report(
    path: Path,
    *,
    pair_count: int,
    factor_count: int,
    case_count: int,
    finding_count: int,
    metrics: Mapping[str, Any],
    holdout_metrics: Mapping[str, Any] | None,
    gate: Mapping[str, Any],
) -> None:
    def percentage(value: Any) -> str:
        if value is None:
            return "N/A"
        return f"{float(value):.1%}"

    lines = [
        "# CarbonFactorResolver 真实数据验收报告",
        "",
        "> 本报告区分闭环摄取一致性验收与独立真实查询 Holdout。提取结果未写入或批准至正式因子库。",
        "",
        "## 验收结论",
        "",
        f"- Release Gate：**{'PASS' if gate['passed'] else 'FAIL'}**",
        f"- 源报告：{pair_count} 组 DOCX/PDF；提取因子：{factor_count} 条",
        f"- 摄取一致性案例：{case_count} 条；来源质量发现：{finding_count} 条",
        "- Error 级来源异常：保留在诊断快照中，`admission_eligible=false`，不得进入正式准入。",
        "",
        "## 闭环摄取一致性验收",
        "",
        "该部分使用同一批提取记录构造隔离 Catalog 与查询，用于验证抽取、证据、准入和闭环检索；不代表未知业务查询的泛化能力。",
        "",
        "| 指标 | 结果 |",
        "|---|---:|",
        f"| Recall@5 | {percentage(metrics.get('recall_at_5'))} |",
        f"| Top-1 | {percentage(metrics.get('top_1_accuracy'))} |",
        f"| 错误候选率 | {percentage(metrics.get('wrong_candidate_rate'))} |",
        f"| 正确拒答率 | {percentage(metrics.get('correct_abstention_rate'))} |",
        f"| MORE_INPUT 正例召回 | {percentage(metrics.get('more_input_positive_recall'))} |",
        f"| MORE_INPUT 负例特异度 | {percentage(metrics.get('more_input_negative_specificity'))} |",
        f"| 证据元数据存在率 | {percentage(metrics.get('evidence_metadata_presence_rate'))} |",
        "",
        "## 独立真实查询 Holdout",
        "",
    ]
    if holdout_metrics is None:
        lines.append("未提供独立 Holdout 清单；Release Gate 不得解释为真实查询能力已通过。")
    else:
        lines.extend([
            "Holdout 查询由静态、人工冻结的业务表达构成，不从运行时 Catalog 名称自动生成。",
            "",
            "| 指标 | 结果 |",
            "|---|---:|",
            f"| 案例数 | {int(holdout_metrics.get('case_count', 0))} |",
            f"| Recall@5 | {percentage(holdout_metrics.get('recall_at_5'))} |",
            f"| Top-1 | {percentage(holdout_metrics.get('top_1_accuracy'))} |",
            f"| 错误候选率 | {percentage(holdout_metrics.get('wrong_candidate_rate'))} |",
            f"| 正确拒答率 | {percentage(holdout_metrics.get('correct_abstention_rate'))} |",
            f"| 拒答负例数 | {int(holdout_metrics.get('abstention_case_count', 0))} |",
            f"| MORE_INPUT 正例召回 | {percentage(holdout_metrics.get('more_input_positive_recall'))} |",
            f"| MORE_INPUT 负例特异度 | {percentage(holdout_metrics.get('more_input_negative_specificity'))} |",
            f"| 证据元数据存在率 | {percentage(holdout_metrics.get('evidence_metadata_presence_rate'))} |",
        ])
    lines.extend([
        "",
        "## Release Gate 明细",
        "",
        "| 检查项 | 结果 |",
        "|---|---:|",
        *(f"| `{name}` | {'PASS' if passed else 'FAIL'} |" for name, passed in gate["checks"].items()),
        "",
        "## 使用限制",
        "",
        "- 本次运行是只读隔离验收，不访问、不修改正式因子数据库或审批存储。",
        "- Release Gate 仅证明当前代码与冻结数据集满足本报告门槛，不等同于因子业务审批。",
        "- 原料与成品、A1/A2/A3/A1-A3、来源质量及准入状态均为硬资格门禁。",
        "",
    ])
    path.write_text("\n".join(lines), encoding="utf-8")


async def run_cases(
    catalog: Mapping[str, Any], cases: list[dict[str, Any]], preset_path: Path
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    before_sha = sha256_file(preset_path)
    database = catalog["database"]

    def fetch_json(_endpoint: str) -> Mapping[str, Any]:
        return catalog

    repository = HttpCatalogFactorRepository(
        endpoint="fixture://true-data-read-only-snapshot",
        expected_sha256=str(database["sha256"]),
        fetch_json=fetch_json,
    )
    engine = A1FactorResolutionEngine(local_retrieval=repository)
    results: list[dict[str, Any]] = []
    for case in cases:
        request = dict(case["request"])
        request["request_id"] = f"true-data:{case['case_id']}"
        started = datetime.now(UTC)
        try:
            recommendation = await engine.resolve(request)
            candidates = (*recommendation.candidates, *recommendation.reviewable_candidates)
            observed = list(dict.fromkeys(candidate.source.source_id for candidate in candidates))
            coverages = [evidence_coverage(candidate) for candidate in candidates[:5]]
            trace = recommendation.trace.explain() if recommendation.trace else None
            error = None
            status = recommendation.status.value
        except Exception as exc:  # a complete acceptance run must retain case failures
            observed, coverages, trace = [], [], None
            error = f"{type(exc).__name__}: {exc}"
            status = ResolutionStatus.ERROR.value
        acceptable = set(case["acceptable_candidates"])
        forbidden = set(case.get("forbidden_candidates", ()))
        observed_top = set(observed[:5])
        wrong_candidates = (
            observed_top - acceptable
            if case.get("forbid_unlisted_candidates")
            else forbidden & observed_top
        )
        results.append(
            {
                "case_id": case["case_id"],
                "case_type": case["case_type"],
                "factor_id": case["factor_id"],
                "expected_more_input": case["expected_more_input"],
                "expected_abstention": case["expected_abstention"],
                "observed_status": status,
                "observed_top_ids": observed[:5],
                "recall_at_5": bool(acceptable & set(observed[:5])) if acceptable else None,
                "recall_at_1": bool(acceptable & set(observed[:1])) if acceptable else None,
                "wrong_candidate_count": len(wrong_candidates),
                "candidate_count": len(observed[:5]),
                "more_input_correct": more_input_is_correct(
                    status, bool(case["expected_more_input"])
                ),
                "abstention_correct": (
                    not observed
                    and status
                    in {
                        ResolutionStatus.UNRESOLVED.value,
                        ResolutionStatus.SUPPLIER_DATA_REQUIRED.value,
                        ResolutionStatus.PROCESS_MODEL_REQUIRED.value,
                    }
                    if case["expected_abstention"]
                    else None
                ),
                "evidence_completeness": sum(coverages) / len(coverages) if coverages else None,
                "evidence_coverage_sum": sum(coverages),
                "evidence_candidate_count": len(coverages),
                "started_at": started.isoformat(),
                "finished_at": datetime.now(UTC).isoformat(),
                "error": error,
                "trace": trace,
            }
        )
    after_sha = sha256_file(preset_path)
    if before_sha != after_sha:
        raise RuntimeError("frozen acceptance preset changed during execution")

    metrics = aggregate_acceptance_metrics(
        results,
        preset_sha_before=before_sha,
        preset_sha_after=after_sha,
        catalog_record_anchor=str(database["sha256"]),
    )
    return results, metrics


def build_acceptance(
    source_dir: Path,
    output_dir: Path,
    *,
    expected_pairs: int | None = None,
    holdout_manifest: Path | None = None,
) -> dict[str, Any]:
    source_dir = source_dir.resolve()
    output_dir = output_dir.resolve()
    if output_dir == source_dir or source_dir in output_dir.parents:
        raise ValueError("output directory must not equal or be inside the source directory")
    if output_dir.exists() and any(output_dir.iterdir()):
        raise ValueError("output directory must be new or empty")
    output_dir.mkdir(parents=True, exist_ok=True)
    pairs = source_pairs(source_dir, expected_pairs=expected_pairs)
    all_factors: list[ExtractedFactor] = []
    all_findings: list[QualityFinding] = []
    for pair in pairs:
        factors, findings = extract_pair(pair)
        all_factors.extend(factors)
        all_findings.extend(findings)
    factors_tuple = tuple(all_factors)
    expected_factor_count = len(pairs) * len(EXPECTED_STAGES)
    if len(factors_tuple) != expected_factor_count:
        raise ValueError(f"expected {expected_factor_count} extracted factors, found {len(factors_tuple)}")

    catalog = catalog_payload(factors_tuple)
    cases = ingestion_cases(factors_tuple)
    manifest_path = output_dir / "ingestion_acceptance_manifest.jsonl"
    write_json(output_dir / "source_manifest.json", [
        {
            "report_id": pair.report_id,
            "product_name_cn": pair.product_name_cn,
            "docx_evidence_id": f"report:{pair.report_id}:docx:{pair.docx_sha256}",
            "pdf_evidence_id": f"report:{pair.report_id}:pdf:{pair.pdf_sha256}",
            "docx_sha256": pair.docx_sha256,
            "pdf_sha256": pair.pdf_sha256,
            "docx_size": pair.docx_size,
            "pdf_size": pair.pdf_size,
        }
        for pair in pairs
    ])
    write_json(output_dir / "extracted_factors.json", [asdict(factor) for factor in factors_tuple])
    write_json(output_dir / "source_quality_findings.json", [asdict(item) for item in all_findings])
    write_json(output_dir / "isolated_catalog_snapshot.json", catalog)
    write_jsonl(manifest_path, cases)
    write_factor_csvs(output_dir, factors_tuple, cases)
    frozen_sha = sha256_file(manifest_path)
    (output_dir / "ingestion_acceptance_manifest.sha256").write_text(
        f"{frozen_sha}  {manifest_path.name}\n", encoding="ascii"
    )
    results, metrics = asyncio.run(run_cases(catalog, cases, manifest_path))
    write_json(output_dir / "cfr_results.json", results)
    write_json(output_dir / "metrics.json", metrics)
    holdout_metrics = None
    holdout_sha = None
    if holdout_manifest is not None:
        holdout_manifest = holdout_manifest.resolve()
        holdout_cases = load_frozen_cases(holdout_manifest)
        holdout_sha = sha256_file(holdout_manifest)
        holdout_results, holdout_metrics = asyncio.run(
            run_cases(catalog, holdout_cases, holdout_manifest)
        )
        write_json(output_dir / "real_query_holdout_results.json", holdout_results)
        write_json(output_dir / "real_query_holdout_metrics.json", holdout_metrics)
    gate = release_gate(metrics, holdout_metrics)
    write_json(output_dir / "release_gate.json", gate)
    write_acceptance_report(
        output_dir / "真实数据验收报告.md",
        pair_count=len(pairs),
        factor_count=len(factors_tuple),
        case_count=len(cases),
        finding_count=len(all_findings),
        metrics=metrics,
        holdout_metrics=holdout_metrics,
        gate=gate,
    )
    try:
        git_sha = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=Path(__file__).resolve().parents[1],
            check=True,
            capture_output=True,
            text=True,
        ).stdout.strip()
    except (OSError, subprocess.CalledProcessError):
        git_sha = None
    try:
        package_version = importlib.metadata.version("carbon-factor-resolver")
    except importlib.metadata.PackageNotFoundError:
        package_version = "unknown"
    source_files_unchanged = all(
        sha256_file(Path(pair.docx_path)) == pair.docx_sha256
        and sha256_file(Path(pair.pdf_path)) == pair.pdf_sha256
        for pair in pairs
    )
    if not source_files_unchanged:
        raise RuntimeError("one or more source reports changed during the acceptance run")
    run_manifest = {
        "schema_version": "cfr-true-data-acceptance-run/v1",
        "generated_at": datetime.now(UTC).isoformat(),
        "git_sha": git_sha,
        "package_version": package_version,
        "source_dataset_id": hashlib.sha256(canonical_json_bytes([
            (pair.report_id, pair.docx_sha256, pair.pdf_sha256) for pair in pairs
        ])).hexdigest(),
        "source_manifest_sha256": sha256_file(output_dir / "source_manifest.json"),
        "extracted_factors_sha256": sha256_file(output_dir / "extracted_factors.json"),
        "catalog_snapshot_sha256": sha256_file(output_dir / "isolated_catalog_snapshot.json"),
        "catalog_record_anchor": metrics["catalog_snapshot_sha256"],
        "ingestion_acceptance_manifest_sha256": frozen_sha,
        "ingestion_acceptance_manifest_verified_unchanged": (
            metrics["preset_sha256_before"] == metrics["preset_sha256_after"]
        ),
        "real_query_holdout_manifest_sha256": holdout_sha,
        "real_query_holdout_metrics": holdout_metrics,
        "release_gate": gate,
        "formal_factor_database_accessed": False,
        "source_files_verified_unchanged": source_files_unchanged,
    }
    write_json(output_dir / "run_manifest.json", run_manifest)
    return {
        "source_pair_count": len(pairs),
        "factor_count": len(factors_tuple),
        "case_count": len(cases),
        "finding_count": len(all_findings),
        "metrics": metrics,
        "real_query_holdout_metrics": holdout_metrics,
        "release_gate": gate,
        "ingestion_manifest_sha256": frozen_sha,
        "output_dir": str(output_dir.resolve()),
    }


def parser() -> argparse.ArgumentParser:
    value = argparse.ArgumentParser(description=__doc__)
    value.add_argument("source_dir", type=Path)
    value.add_argument("output_dir", type=Path)
    value.add_argument("--expected-pairs", type=int)
    value.add_argument("--holdout-manifest", type=Path)
    return value


def main() -> None:
    args = parser().parse_args()
    summary = build_acceptance(
        args.source_dir.resolve(),
        args.output_dir.resolve(),
        expected_pairs=args.expected_pairs,
        holdout_manifest=args.holdout_manifest,
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    if not summary["release_gate"]["passed"]:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
